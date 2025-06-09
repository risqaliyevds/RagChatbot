# SPDX-License-Identifier: Apache-2.0
"""
FastAPI RAG Application with Qdrant Vector Store and PostgreSQL Database
========================================================================

Enhanced version with PostgreSQL database for chat history storage
instead of JSON files, with user management and chat history tracking
"""

import os
import logging
import json
import re
from typing import List, Optional, Dict, Any, Tuple
from contextlib import asynccontextmanager
import uuid
import time
from datetime import datetime, timedelta
from pathlib import Path

from fastapi import FastAPI, HTTPException, status, UploadFile, File, Form
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, Field
from dotenv import load_dotenv
from fastapi.middleware.cors import CORSMiddleware

# LangChain imports
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import WebBaseLoader, DirectoryLoader, TextLoader, PyPDFLoader, UnstructuredWordDocumentLoader
from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI, OpenAIEmbeddings

# Qdrant imports
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct, Filter, FieldCondition, MatchValue

# Transformers imports for embedding
import torch
from transformers import AutoTokenizer, AutoModel
import numpy as np

# Database imports
from database import DatabaseManager, get_db_manager, init_database

# Try to import HuggingFaceEmbeddings for fallback support
HUGGINGFACE_AVAILABLE = False
try:
    from langchain_community.embeddings import HuggingFaceEmbeddings
    HUGGINGFACE_AVAILABLE = True
except ImportError:
    pass

# Load environment variables
load_dotenv(".env")  # Load .env first if it exists
load_dotenv("config.env")  # Then load config.env (will override .env values)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Global variables for RAG components
vectorstore = None
qa_chain = None
qdrant_client = None
embeddings_model = None
db_manager = None

# Chat history storage - now using PostgreSQL
CHAT_HISTORY_FILE = "chat_history.json"  # For migration purposes
CHAT_TIMEOUT_HOURS = 1

# Add a new import for progress tracking
import asyncio
from typing import AsyncGenerator


class MultilingualE5Embeddings:
    """Custom embedding class for intfloat/multilingual-e5-large-instruct model"""
    
    def __init__(self, model_name: str = "intfloat/multilingual-e5-large-instruct", device: str = None):
        self.model_name = model_name
        self.device = device or ("cuda" if torch.cuda.is_available() else "cpu")
        
        logger.info(f"Loading multilingual E5 model: {model_name} on {self.device}")
        
        # Load tokenizer and model
        self.tokenizer = AutoTokenizer.from_pretrained(model_name)
        self.model = AutoModel.from_pretrained(model_name)
        self.model.to(self.device)
        self.model.eval()
        
        logger.info(f"Successfully loaded {model_name}")
    
    def _get_detailed_instruct(self, task_description: str, query: str) -> str:
        """Format query with task description for E5 instruct model"""
        return f"Instruct: {task_description}\nQuery: {query}"
    
    def _encode_text(self, texts: List[str], task_description: str = "Given a question, retrieve passages that answer the question") -> np.ndarray:
        """Encode texts using the E5 model"""
        # Format texts with instruction
        formatted_texts = [self._get_detailed_instruct(task_description, text) for text in texts]
        
        # Tokenize
        batch_dict = self.tokenizer(
            formatted_texts,
            max_length=512,
            padding=True,
            truncation=True,
            return_tensors='pt'
        )
        
        # Move to device
        batch_dict = {k: v.to(self.device) for k, v in batch_dict.items()}
        
        # Get embeddings
        with torch.no_grad():
            outputs = self.model(**batch_dict)
            embeddings = self._average_pool(outputs.last_hidden_state, batch_dict['attention_mask'])
            
        # Normalize embeddings
        embeddings = torch.nn.functional.normalize(embeddings, p=2, dim=1)
        
        return embeddings.cpu().numpy()
    
    def _average_pool(self, last_hidden_states: torch.Tensor, attention_mask: torch.Tensor) -> torch.Tensor:
        """Average pooling with attention mask"""
        last_hidden = last_hidden_states.masked_fill(~attention_mask[..., None].bool(), 0.0)
        return last_hidden.sum(dim=1) / attention_mask.sum(dim=1)[..., None]
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Embed a list of documents"""
        embeddings = self._encode_text(texts, "Given a question, retrieve passages that answer the question")
        return embeddings.tolist()
    
    def embed_query(self, text: str) -> List[float]:
        """Embed a single query"""
        embeddings = self._encode_text([text], "Given a question, retrieve passages that answer the question")
        return embeddings[0].tolist()


class ChatMessage(BaseModel):
    """Ð§Ð°Ñ‚ Ñ…Ð°Ð±Ð°Ñ€Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    role: str = Field(..., description="Ð¥Ð°Ð±Ð°Ñ€ ÑŽÐ±Ð¾Ñ€ÑƒÐ²Ñ‡Ð¸Ð½Ð¸Ð½Ð³ Ñ€Ð¾Ð»Ð¸")
    content: str = Field(..., description="Ð¥Ð°Ð±Ð°Ñ€ Ð¼Ð°Ð·Ð¼ÑƒÐ½Ð¸")
    timestamp: Optional[datetime] = Field(default_factory=datetime.now, description="Ð¥Ð°Ð±Ð°Ñ€ Ð²Ð°Ò›Ñ‚Ð¸")


class ChatSession(BaseModel):
    """Ð§Ð°Ñ‚ ÑÐµÑÑÐ¸ÑÑÐ¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    chat_id: str = Field(..., description="Ð§Ð°Ñ‚ Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€Ð¸")
    user_id: str = Field(..., description="Ð¤Ð¾Ð¹Ð´Ð°Ð»Ð°Ð½ÑƒÐ²Ñ‡Ð¸ Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€Ð¸")
    messages: List[ChatMessage] = Field(default_factory=list, description="Ð§Ð°Ñ‚ Ñ…Ð°Ð±Ð°Ñ€Ð»Ð°Ñ€Ð¸")
    created_at: datetime = Field(default_factory=datetime.now, description="Ð¯Ñ€Ð°Ñ‚Ð¸Ð»Ð³Ð°Ð½ Ð²Ð°Ò›Ñ‚")
    last_activity: datetime = Field(default_factory=datetime.now, description="Ð¡ÑžÐ½Ð³Ð³Ð¸ Ñ„Ð°Ð¾Ð»Ð¸ÑÑ‚")


class ChatRequest(BaseModel):
    """Ð§Ð°Ñ‚ ÑÑžÑ€Ð¾Ð²Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    user_id: str = Field(..., description="Ð¤Ð¾Ð¹Ð´Ð°Ð»Ð°Ð½ÑƒÐ²Ñ‡Ð¸ Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€Ð¸")
    chat_id: Optional[str] = Field(None, description="Ð§Ð°Ñ‚ Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€Ð¸ (ÑÐ½Ð³Ð¸ Ñ‡Ð°Ñ‚ ÑƒÑ‡ÑƒÐ½ Ð±ÑžÑˆ Ò›Ð¾Ð»Ð´Ð¸Ñ€Ð¸Ð½Ð³)")
    message: str = Field(..., description="Ð¤Ð¾Ð¹Ð´Ð°Ð»Ð°Ð½ÑƒÐ²Ñ‡Ð¸ Ñ…Ð°Ð±Ð°Ñ€Ð¸")


class ChatResponse(BaseModel):
    """Ð§Ð°Ñ‚ Ð¶Ð°Ð²Ð¾Ð±Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    chat_id: str = Field(..., description="Ð§Ð°Ñ‚ Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€Ð¸")
    user_id: str = Field(..., description="Ð¤Ð¾Ð¹Ð´Ð°Ð»Ð°Ð½ÑƒÐ²Ñ‡Ð¸ Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€Ð¸")
    message: str = Field(..., description="Ð‘Ð¾Ñ‚ Ð¶Ð°Ð²Ð¾Ð±Ð¸")
    timestamp: datetime = Field(default_factory=datetime.now, description="Ð–Ð°Ð²Ð¾Ð± Ð²Ð°Ò›Ñ‚Ð¸")


class ChatHistoryRequest(BaseModel):
    """Ð§Ð°Ñ‚ Ñ‚Ð°Ñ€Ð¸Ñ…Ð¸ ÑÑžÑ€Ð¾Ð²Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    user_id: str = Field(..., description="Ð¤Ð¾Ð¹Ð´Ð°Ð»Ð°Ð½ÑƒÐ²Ñ‡Ð¸ Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€Ð¸")
    chat_id: str = Field(..., description="Ð§Ð°Ñ‚ Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€Ð¸")


class ChatCompletionRequest(BaseModel):
    """Ð§Ð°Ñ‚ Ñ‚ÑžÐ»Ð´Ð¸Ñ€Ð¸Ñˆ ÑÑžÑ€Ð¾Ð²Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    model: str = Field(default="google/gemma-3-12b-it", description="Ð˜ÑˆÐ»Ð°Ñ‚Ð¸Ð»Ð°Ð´Ð¸Ð³Ð°Ð½ Ð¼Ð¾Ð´ÐµÐ»")
    messages: List[ChatMessage] = Field(..., description="Ð¥Ð°Ð±Ð°Ñ€Ð»Ð°Ñ€ Ñ€ÑžÐ¹Ñ…Ð°Ñ‚Ð¸")
    max_tokens: Optional[int] = Field(default=150, description="Ð˜ÑˆÐ»Ð°Ð± Ñ‡Ð¸Ò›Ð°Ñ€Ð¸Ð»Ð°Ð´Ð¸Ð³Ð°Ð½ Ð¼Ð°ÐºÑÐ¸Ð¼Ð°Ð» Ñ‚Ð¾ÐºÐµÐ½Ð»Ð°Ñ€")
    temperature: Optional[float] = Field(default=0.7, description="ÐÐ°Ð¼ÑƒÐ½Ð° Ð¾Ð»Ð¸Ñˆ Ò³Ð°Ñ€Ð¾Ñ€Ð°Ñ‚Ð¸")
    stream: Optional[bool] = Field(default=False, description="Ð–Ð°Ð²Ð¾Ð±Ð»Ð°Ñ€Ð½Ð¸ Ð¾Ò›Ð¸Ð¼ Ñ‚Ð°Ñ€Ð·Ð¸Ð´Ð° ÑŽÐ±Ð¾Ñ€Ð¸Ñˆ")


class ChatCompletionResponse(BaseModel):
    """Ð§Ð°Ñ‚ Ñ‚ÑžÐ»Ð´Ð¸Ñ€Ð¸Ñˆ Ð¶Ð°Ð²Ð¾Ð±Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    id: str = Field(..., description="Ð¢ÑžÐ»Ð´Ð¸Ñ€Ð¸Ñˆ ÑƒÑ‡ÑƒÐ½ Ð½Ð¾Ñ‘Ð± Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€")
    object: str = Field(default="chat.completion", description="ÐžÐ±ÑŠÐµÐºÑ‚ Ñ‚ÑƒÑ€Ð¸")
    created: int = Field(..., description="Ð¯Ñ€Ð°Ñ‚Ð¸Ð»Ð³Ð°Ð½ Unix Ð²Ð°Ò›Ñ‚ Ð±ÐµÐ»Ð³Ð¸ÑÐ¸")
    model: str = Field(..., description="Ð¢ÑžÐ»Ð´Ð¸Ñ€Ð¸Ñˆ ÑƒÑ‡ÑƒÐ½ Ð¸ÑˆÐ»Ð°Ñ‚Ð¸Ð»Ð³Ð°Ð½ Ð¼Ð¾Ð´ÐµÐ»")
    choices: List[Dict[str, Any]] = Field(..., description="Ð¢ÑžÐ»Ð´Ð¸Ñ€Ð¸Ñˆ Ñ‚Ð°Ð½Ð»Ð¾Ð²Ð»Ð°Ñ€Ð¸ Ñ€ÑžÐ¹Ñ…Ð°Ñ‚Ð¸")
    usage: Dict[str, int] = Field(..., description="Ð¢Ð¾ÐºÐµÐ½ Ð¸ÑˆÐ»Ð°Ñ‚Ð¸Ñˆ Ð¼Ð°ÑŠÐ»ÑƒÐ¼Ð¾Ñ‚Ð¸")


class NewChatRequest(BaseModel):
    """Ð¯Ð½Ð³Ð¸ Ñ‡Ð°Ñ‚ ÑÑ€Ð°Ñ‚Ð¸Ñˆ ÑÑžÑ€Ð¾Ð²Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    user_id: str = Field(..., description="Ð¤Ð¾Ð¹Ð´Ð°Ð»Ð°Ð½ÑƒÐ²Ñ‡Ð¸ Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€Ð¸")


class NewChatResponse(BaseModel):
    """Ð¯Ð½Ð³Ð¸ Ñ‡Ð°Ñ‚ ÑÑ€Ð°Ñ‚Ð¸Ñˆ Ð¶Ð°Ð²Ð¾Ð±Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    chat_id: str = Field(..., description="Ð¯Ð½Ð³Ð¸ Ñ‡Ð°Ñ‚ Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€Ð¸")
    user_id: str = Field(..., description="Ð¤Ð¾Ð¹Ð´Ð°Ð»Ð°Ð½ÑƒÐ²Ñ‡Ð¸ Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€Ð¸")
    message: str = Field(..., description="Ð¥ÑƒÑˆ ÐºÐµÐ»Ð¸Ð±ÑÐ¸Ð· Ñ…Ð°Ð±Ð°Ñ€Ð¸")
    created_at: str = Field(..., description="Ð¯Ñ€Ð°Ñ‚Ð¸Ð»Ð³Ð°Ð½ Ð²Ð°Ò›Ñ‚")
    last_activity: str = Field(..., description="Ð¡ÑžÐ½Ð³Ð³Ð¸ Ñ„Ð°Ð¾Ð»Ð¸ÑÑ‚")


class UserSessionStatusRequest(BaseModel):
    """Ð¤Ð¾Ð¹Ð´Ð°Ð»Ð°Ð½ÑƒÐ²Ñ‡Ð¸ ÑÐµÑÑÐ¸Ñ Ò³Ð¾Ð»Ð°Ñ‚Ð¸ ÑÑžÑ€Ð¾Ð²Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    user_id: str = Field(..., description="Ð¤Ð¾Ð¹Ð´Ð°Ð»Ð°Ð½ÑƒvÑ‡Ð¸ Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€Ð¸")


class UserSessionStatusResponse(BaseModel):
    """Ð¤Ð¾Ð¹Ð´Ð°Ð»Ð°Ð½ÑƒÐ²Ñ‡Ð¸ ÑÐµÑÑÐ¸Ñ Ò³Ð¾Ð»Ð°Ñ‚Ð¸ Ð¶Ð°Ð²Ð¾Ð±Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    user_id: str = Field(..., description="Ð¤Ð¾Ð¹Ð´Ð°Ð»Ð°Ð½ÑƒÐ²Ñ‡Ð¸ Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€Ð¸")
    has_active_session: bool = Field(..., description="Ð¤Ð°Ð¾Ð» ÑÐµÑÑÐ¸Ñ Ð¼Ð°Ð²Ð¶ÑƒÐ´Ð»Ð¸Ð³Ð¸")
    active_chat_id: Optional[str] = Field(None, description="Ð¤Ð°Ð¾Ð» Ñ‡Ð°Ñ‚ Ð¸Ð´ÐµÐ½Ñ‚Ð¸Ñ„Ð¸ÐºÐ°Ñ‚Ð¾Ñ€Ð¸")
    last_activity: Optional[str] = Field(None, description="Ð¡ÑžÐ½Ð³Ð³Ð¸ Ñ„Ð°Ð¾Ð»Ð¸ÑÑ‚ Ð²Ð°Ò›Ñ‚Ð¸")
    session_expired: bool = Field(..., description="Ð¡ÐµÑÑÐ¸Ñ Ð¼ÑƒddaÑ‚Ð¸ Ñ‚ÑƒÐ³Ð°Ð³Ð°Ð½Ð¼Ð¸")


class HealthResponse(BaseModel):
    """Ð¡Ð¾Ò“Ð»Ð¸Ò›Ð½Ð¸ Ñ‚ÐµÐºÑˆÐ¸Ñ€Ð¸Ñˆ Ð¶Ð°Ð²Ð¾Ð± Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    status: str = Field(..., description="Ð¡Ð¾Ò“Ð»Ð¸Ò› Ò³Ð¾Ð»Ð°Ñ‚Ð¸")
    message: str = Field(..., description="Ð¡Ð¾Ò“Ð»Ð¸Ò› Ñ…Ð°Ð±Ð°Ñ€Ð¸")
    qdrant_status: str = Field(..., description="Qdrant ÑƒÐ»Ð°Ð½Ð¸Ñˆ Ò³Ð¾Ð»Ð°Ñ‚Ð¸")
    database_status: str = Field(..., description="Database ÑƒÐ»Ð°Ð½Ð¸Ñˆ Ò³Ð¾Ð»Ð°Ñ‚Ð¸")


class DocumentUploadResponse(BaseModel):
    """Ò²ÑƒÐ¶Ð¶Ð°Ñ‚ ÑŽÐºÐ»Ð°Ñˆ Ð¶Ð°Ð²Ð¾Ð±Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    success: bool = Field(..., description="Ð®ÐºÐ»Ð°Ñˆ Ð¼ÑƒÐ²Ð°Ñ„Ñ„Ð°Ò›Ð¸ÑÑ‚Ð»Ð¸ Ð±ÑžÐ»Ð´Ð¸Ð¼Ð¸")
    message: str = Field(..., description="Ð–Ð°Ð²Ð¾Ð± Ñ…Ð°Ð±Ð°Ñ€Ð¸")
    filename: str = Field(..., description="Ð®ÐºÐ»Ð°Ð½Ð³Ð°Ð½ Ñ„Ð°Ð¹Ð» Ð½Ð¾Ð¼Ð¸")
    file_size: int = Field(..., description="Ð¤Ð°Ð¹Ð» Ò³Ð°Ð¶Ð¼Ð¸ (Ð±Ð°Ð¹Ñ‚Ð»Ð°Ñ€Ð´Ð°)")
    chunks_added: int = Field(..., description="ÒšÑžÑˆÐ¸Ð»Ð³Ð°Ð½ Ñ‡Ð°Ð½ÐºÐ»Ð°Ñ€ ÑÐ¾Ð½Ð¸")
    processing_time: float = Field(..., description="Ð˜ÑˆÐ»Ð¾Ð² Ð±ÐµÑ€Ð¸Ñˆ Ð²Ð°Ò›Ñ‚Ð¸ (ÑÐ¾Ð½Ð¸ÑÐ»Ð°Ñ€Ð´Ð°)")


class FileInfo(BaseModel):
    """Ð¤Ð°Ð¹Ð» Ð¼Ð°ÑŠÐ»ÑƒÐ¼Ð¾Ñ‚Ð»Ð°Ñ€Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    filename: str = Field(..., description="Ð¤Ð°Ð¹Ð» Ð½Ð¾Ð¼Ð¸")
    file_size: int = Field(..., description="Ð¤Ð°Ð¹Ð» Ò³Ð°Ð¶Ð¼Ð¸ (Ð±Ð°Ð¹Ñ‚Ð»Ð°Ñ€Ð´Ð°)")
    created_at: str = Field(..., description="Ð¯Ñ€Ð°Ñ‚Ð¸Ð»Ð³Ð°Ð½ ÑÐ°Ð½Ð°")
    modified_at: str = Field(..., description="ÐŽÐ·Ð³Ð°Ñ€Ñ‚Ð¸Ñ€Ð¸Ð»Ð³Ð°Ð½ ÑÐ°Ð½Ð°")
    file_extension: str = Field(..., description="Ð¤Ð°Ð¹Ð» ÐºÐµÐ½Ð³Ð°Ð¹Ñ‚Ð¼Ð°ÑÐ¸")


class DocumentListResponse(BaseModel):
    """Ò²ÑƒÐ¶Ð¶Ð°Ñ‚Ð»Ð°Ñ€ Ñ€ÑžÐ¹Ñ…Ð°Ñ‚Ð¸ Ð¶Ð°Ð²Ð¾Ð±Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    success: bool = Field(..., description="ÐÐ¼Ð°Ð» Ð¼ÑƒÐ²Ð°Ñ„Ñ„Ð°Ò›Ð¸ÑÑ‚Ð»Ð¸ Ð±ÑžÐ»Ð´Ð¸Ð¼Ð¸")
    message: str = Field(..., description="Ð–Ð°Ð²Ð¾Ð± Ñ…Ð°Ð±Ð°Ñ€Ð¸")
    files: List[FileInfo] = Field(default_factory=list, description="Ð¤Ð°Ð¹Ð»Ð»Ð°Ñ€ Ñ€ÑžÐ¹Ñ…Ð°Ñ‚Ð¸")
    total_files: int = Field(..., description="Ð–Ð°Ð¼Ð¸ Ñ„Ð°Ð¹Ð»Ð»Ð°Ñ€ ÑÐ¾Ð½Ð¸")
    total_size: int = Field(..., description="Ð–Ð°Ð¼Ð¸ Ñ„Ð°Ð¹Ð»Ð»Ð°Ñ€ Ò³Ð°Ð¶Ð¼Ð¸")


class DocumentDeleteRequest(BaseModel):
    """Ò²ÑƒÐ¶Ð¶Ð°Ñ‚ ÑžÑ‡Ð¸Ñ€Ð¸Ñˆ ÑÑžÑ€Ð¾Ð²Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    filename: str = Field(..., description="ÐŽÑ‡Ð¸Ñ€Ð¸Ð»Ð¸ÑˆÐ¸ ÐºÐµÑ€Ð°Ðº Ð±ÑžÐ»Ð³Ð°Ð½ Ñ„Ð°Ð¹Ð» Ð½Ð¾Ð¼Ð¸")


class DocumentDeleteResponse(BaseModel):
    """Ò²ÑƒÐ¶Ð¶Ð°Ñ‚ ÑžÑ‡Ð¸Ñ€Ð¸Ñˆ Ð¶Ð°Ð²Ð¾Ð±Ð¸ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    success: bool = Field(..., description="ÐŽÑ‡Ð¸Ñ€Ð¸Ñˆ Ð¼ÑƒÐ²Ð°Ñ„Ñ„Ð°Ò›Ð¸ÑÑ‚Ð»Ð¸ Ð±ÑžÐ»Ð´Ð¸Ð¼Ð¸")
    message: str = Field(..., description="Ð–Ð°Ð²Ð¾Ð± Ñ…Ð°Ð±Ð°Ñ€Ð¸")
    filename: str = Field(..., description="ÐŽÑ‡Ð¸Ñ€Ð¸Ð»Ð³Ð°Ð½ Ñ„Ð°Ð¹Ð» Ð½Ð¾Ð¼Ð¸")
    embeddings_deleted: int = Field(default=0, description="ÐŽÑ‡Ð¸Ñ€Ð¸Ð»Ð³Ð°Ð½ Ð²ÐµÐºÑ‚Ð¾Ñ€Ð»Ð°Ñ€ ÑÐ¾Ð½Ð¸")





class DocumentUploadProgress(BaseModel):
    """Ò²ÑƒÐ¶Ð¶Ð°Ñ‚ ÑŽÐºÐ»Ð°Ñˆ Ð¿Ñ€Ð¾Ð³Ñ€ÐµÑÑ Ð¼Ð¾Ð´ÐµÐ»Ð¸"""
    stage: str = Field(..., description="Ò²Ð¾Ð·Ð¸Ñ€Ð³Ð¸ Ð±Ð¾ÑÒ›Ð¸Ñ‡")
    progress: float = Field(..., description="ÐŸÑ€Ð¾Ð³Ñ€ÐµÑÑ Ñ„Ð¾Ð¸Ð·Ð¸ (0-100)")
    message: str = Field(..., description="Ò²Ð¾Ð·Ð¸Ñ€Ð³Ð¸ Ò³Ð¾Ð»Ð°Ñ‚ Ñ…Ð°Ð±Ð°Ñ€Ð¸")
    details: Optional[Dict[str, Any]] = Field(default=None, description="ÒšÑžÑˆÐ¸Ð¼Ñ‡Ð° Ð¼Ð°ÑŠÐ»ÑƒÐ¼Ð¾Ñ‚Ð»Ð°Ñ€")


def load_config() -> Dict[str, Any]:
    """ÐœÑƒÒ³Ð¸Ñ‚ ÑžÐ·Ð³Ð°Ñ€ÑƒÐ²Ñ‡Ð¸Ð»Ð°Ñ€Ð¸Ð´Ð°Ð½ ÐºÐ¾Ð½Ñ„Ð¸Ð³ÑƒÑ€Ð°Ñ†Ð¸ÑÐ½Ð¸ ÑŽÐºÐ»Ð°Ñˆ"""
    embedding_port = os.getenv("EMBEDDING_PORT", "8000")
    chat_port = os.getenv("CHAT_MODEL_PORT", "8001")
    
    return {
        "vllm_api_key": os.getenv("VLLM_API_KEY", "EMPTY"),
        "vllm_embedding_endpoint": os.getenv("VLLM_EMBEDDING_ENDPOINT"),  # Can be None
        "vllm_chat_endpoint": os.getenv("VLLM_CHAT_ENDPOINT", f"http://localhost:{chat_port}/v1"),
        # Database configuration
        "database_url": os.getenv("DATABASE_URL", "postgresql://chatbot_user:chatbot_password@localhost:5432/chatbot_db"),
        # Qdrant configuration
        "qdrant_url": os.getenv("QDRANT_URL", None),
        "qdrant_api_key": os.getenv("QDRANT_API_KEY", None),
        "qdrant_path": os.getenv("QDRANT_PATH", "./qdrant_storage"),
        "qdrant_collection_name": os.getenv("QDRANT_COLLECTION_NAME", "rag_documents"),
        # Qdrant Vector Configuration
        "qdrant_vector_size": int(os.getenv("QDRANT_VECTOR_SIZE", "0")),  # 0 means auto-detect
        "qdrant_distance": os.getenv("QDRANT_DISTANCE", "COSINE").upper(),
        "qdrant_force_recreate": os.getenv("QDRANT_FORCE_RECREATE", "false").lower() == "true",
        "qdrant_on_disk": os.getenv("QDRANT_ON_DISK", "false").lower() == "true",
        # Document configuration
        "document_url": os.getenv("DOCUMENT_URL", "https://docs.vllm.ai/en/latest/getting_started/quickstart.html"),
        "documents_path": os.getenv("DOCUMENTS_PATH", "/app/documents"),
        # Model configuration
        "embedding_model": os.getenv("EMBEDDING_MODEL", "intfloat/multilingual-e5-large-instruct"),
        "chat_model": os.getenv("CHAT_MODEL", "Qwen/Qwen3-8B"),
        # RAG parameters
        "top_k": int(os.getenv("TOP_K", "5")),
        "chunk_size": int(os.getenv("CHUNK_SIZE", "1000")),
        "chunk_overlap": int(os.getenv("CHUNK_OVERLAP", "200")),
        # Server configuration
        "host": os.getenv("HOST", "0.0.0.0"),
        "port": int(os.getenv("PORT", "8080")),
    }


def init_qdrant_client(config: Dict[str, Any]) -> QdrantClient:
    """Qdrant ÐºÐ»Ð¸ÐµÐ½Ñ‚Ð¸Ð½Ð¸ Ð¸ÑˆÐ³Ð° Ñ‚ÑƒÑˆÐ¸Ñ€Ð¸Ñˆ"""
    if config["qdrant_url"]:
        logger.info(f"Connecting to Qdrant at: {config['qdrant_url']}")
        return QdrantClient(
            url=config["qdrant_url"],
            api_key=config["qdrant_api_key"]
        )
    else:
        logger.info(f"Using local Qdrant storage at: {config['qdrant_path']}")
        os.makedirs(config["qdrant_path"], exist_ok=True)
        return QdrantClient(path=config["qdrant_path"])


def load_and_split_documents(config: Dict[str, Any]) -> List[Document]:
    """Ð’ÐµÐ± URL Ñ‘ÐºÐ¸ Ð¼Ð°Ñ…Ð°Ð»Ð»Ð¸Ð¹ Ð´Ð¸Ñ€ÐµÐºÑ‚Ð¾Ñ€Ð¸ÑÐ´Ð°Ð½ Ò³ÑƒÐ¶Ð¶Ð°Ñ‚Ð»Ð°Ñ€Ð½Ð¸ ÑŽÐºÐ»Ð°Ñˆ Ð²Ð° Ð±ÑžÐ»Ð¸Ñˆ"""
    try:
        documents = []
        
        # Check if local documents directory exists
        if os.path.exists(config["documents_path"]) and os.path.isdir(config["documents_path"]):
            logger.info(f"Loading documents from local directory: {config['documents_path']}")
            
            # Load text files
            try:
                text_loader = DirectoryLoader(
                    config["documents_path"],
                    glob="**/*.txt",
                    loader_cls=TextLoader,
                    show_progress=True
                )
                docs = text_loader.load()
                if docs:
                    logger.info(f"Loaded {len(docs)} text files")
                    documents.extend(docs)
            except Exception as e:
                logger.warning(f"Could not load text files: {e}")
            
            # Load Word documents
            try:
                import glob
                docx_pattern = os.path.join(config["documents_path"], "**/*.docx")
                docx_files = glob.glob(docx_pattern, recursive=True)
                
                for docx_file in docx_files:
                    try:
                        docx_loader = UnstructuredWordDocumentLoader(docx_file)
                        docx_docs = docx_loader.load()
                        documents.extend(docx_docs)
                        logger.info(f"Loaded Word document: {docx_file}")
                    except Exception as e:
                        logger.warning(f"Could not load Word document {docx_file}: {e}")
            except Exception as e:
                logger.warning(f"Could not process Word documents: {e}")
            
            # Load PDF files
            try:
                import glob
                pdf_pattern = os.path.join(config["documents_path"], "**/*.pdf")
                pdf_files = glob.glob(pdf_pattern, recursive=True)
                
                for pdf_file in pdf_files:
                    try:
                        pdf_loader = PyPDFLoader(pdf_file)
                        pdf_docs = pdf_loader.load()
                        documents.extend(pdf_docs)
                        logger.info(f"Loaded PDF document: {pdf_file}")
                    except Exception as e:
                        logger.warning(f"Could not load PDF document {pdf_file}: {e}")
            except Exception as e:
                logger.warning(f"Could not process PDF documents: {e}")
            
            # Load markdown files
            try:
                md_loader = DirectoryLoader(
                    config["documents_path"],
                    glob="**/*.md",
                    loader_cls=TextLoader,
                    show_progress=True
                )
                md_docs = md_loader.load()
                if md_docs:
                    logger.info(f"Loaded {len(md_docs)} markdown files")
                    documents.extend(md_docs)
            except Exception as e:
                logger.warning(f"Could not load markdown files: {e}")
            
            # Load Python files
            try:
                py_loader = DirectoryLoader(
                    config["documents_path"],
                    glob="**/*.py",
                    loader_cls=TextLoader,
                    show_progress=True
                )
                py_docs = py_loader.load()
                if py_docs:
                    logger.info(f"Loaded {len(py_docs)} Python files")
                    documents.extend(py_docs)
            except Exception as e:
                logger.warning(f"Could not load Python files: {e}")
        
        # If no local documents found, try web URL
        if not documents:
            logger.info(f"No local documents found, loading from web URL: {config['document_url']}")
            try:
                loader = WebBaseLoader(config["document_url"])
                documents = loader.load()
                logger.info(f"Loaded {len(documents)} documents from web")
            except Exception as e:
                logger.error(f"Failed to load documents from web: {e}")
                # Create a fallback document
                documents = [Document(
                    page_content="This is a fallback document for testing purposes.",
                    metadata={"source": "fallback", "title": "Test Document"}
                )]
        
        if not documents:
            raise ValueError("No documents could be loaded")
        
        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config["chunk_size"],
            chunk_overlap=config["chunk_overlap"],
            length_function=len,
        )
        
        split_docs = text_splitter.split_documents(documents)
        logger.info(f"Split {len(documents)} documents into {len(split_docs)} chunks")
        
        return split_docs
        
    except Exception as e:
        logger.error(f"Error loading documents: {e}")
        raise


class SimpleQdrantVectorStore:
    """Simplified Qdrant vector store implementation"""
    
    def __init__(self, client: QdrantClient, collection_name: str, embeddings, config: Dict[str, Any]):
        self.client = client
        self.collection_name = collection_name
        self.embeddings = embeddings
        self.config = config
    
    async def add_documents(self, documents: List[Document], embedding_dim: int = None, progress_callback=None):
        """Add documents to the vector store"""
        try:
            # Get embedding dimension from first document if not provided
            if embedding_dim is None:
                test_embedding = self.embeddings.embed_query("test")
                embedding_dim = len(test_embedding)
                logger.info(f"Auto-detected embedding dimension: {embedding_dim}")
            
            # Check if collection exists
            try:
                collection_info = self.client.get_collection(self.collection_name)
                existing_dim = collection_info.config.params.vectors.size
                
                if existing_dim != embedding_dim:
                    if self.config.get("qdrant_force_recreate", False):
                        logger.warning(f"Dimension mismatch ({existing_dim} vs {embedding_dim}). Recreating collection.")
                        self.client.delete_collection(self.collection_name)
                        raise Exception("Force recreate")
                    else:
                        raise ValueError(f"Embedding dimension mismatch: existing={existing_dim}, new={embedding_dim}")
                        
                logger.info(f"Using existing collection: {self.collection_name}")
                
            except Exception as e:
                # Create new collection
                logger.info(f"Creating new collection: {self.collection_name}")
                
                # Map distance metric
                distance_map = {
                    "COSINE": Distance.COSINE,
                    "EUCLID": Distance.EUCLID,
                    "DOT": Distance.DOT
                }
                distance = distance_map.get(self.config.get("qdrant_distance", "COSINE"), Distance.COSINE)
                
                self.client.create_collection(
                    collection_name=self.collection_name,
                    vectors_config=VectorParams(
                        size=embedding_dim,
                        distance=distance,
                        on_disk=self.config.get("qdrant_on_disk", False)
                    )
                )
            
            # Always add new documents (removed the early return bug)
            collection_info = self.client.get_collection(self.collection_name)
            existing_count = collection_info.points_count
            logger.info(f"Collection currently contains {existing_count} documents. Adding {len(documents)} new documents.")
            
            # Debug: Check if documents list is empty
            if not documents:
                logger.warning("âš ï¸ No documents to add - documents list is empty!")
                return
            
            # Prepare documents for embedding
            texts = [doc.page_content for doc in documents]
            metadatas = [doc.metadata for doc in documents]
            
            logger.info(f"Embedding {len(texts)} documents...")
            
            # Generate embeddings in batches - increased batch size for better performance
            batch_size = 50  # Larger batches for faster processing while avoiding memory issues
            all_embeddings = []
            total_batches = (len(texts) + batch_size - 1) // batch_size
            
            for i in range(0, len(texts), batch_size):
                batch_texts = texts[i:i + batch_size]
                batch_embeddings = self.embeddings.embed_documents(batch_texts)
                all_embeddings.extend(batch_embeddings)
                
                current_batch = i//batch_size + 1
                progress_percent = (current_batch / total_batches) * 100
                logger.info(f"ðŸ“Š Embedding progress: {current_batch}/{total_batches} batches ({progress_percent:.1f}%)")
                
                # Store progress for SSE endpoint access
                if progress_callback:
                    await progress_callback(f"embedding_batch_{current_batch}", 30 + (progress_percent * 0.4), 
                                          f"Embedding batch {current_batch}/{total_batches}")
                
                # Add small delay for very fast processing to allow progress updates
                if total_batches > 10:
                    await asyncio.sleep(0.1)
            
            # Create points for Qdrant - use UUID-based IDs for reliability
            points = []
            point_ids = []
            chunks_data = []
            
            for i, (text, metadata, embedding) in enumerate(zip(texts, metadatas, all_embeddings)):
                # Generate unique point ID using UUID
                import time
                point_id = int(str(int(time.time() * 1000))[-10:] + str(i).zfill(6))  # Timestamp + index
                
                # Ensure metadata is JSON serializable
                clean_metadata = {}
                for key, value in metadata.items():
                    if isinstance(value, (str, int, float, bool, list, dict)):
                        clean_metadata[key] = value
                    else:
                        clean_metadata[key] = str(value)
                
                # Add text content to metadata for retrieval
                clean_metadata["content"] = text[:1000]  # Limit content length
                
                point = PointStruct(
                    id=point_id,
                    vector=embedding,
                    payload=clean_metadata
                )
                points.append(point)
                point_ids.append(point_id)
                
                # Store chunk data for database tracking
                chunks_data.append({
                    "content": text,
                    "metadata": clean_metadata
                })
            
            # Upload points to Qdrant
            logger.info(f"Uploading {len(points)} points to Qdrant...")
            try:
                self.client.upsert(
                    collection_name=self.collection_name,
                    points=points
                )
                logger.info(f"âœ… Successfully uploaded {len(points)} points to Qdrant")
                
                # Store point IDs in database for tracking
                try:
                    # Extract filename from first document metadata
                    filename = None
                    file_size = 0
                    for doc in documents:
                        if 'uploaded_filename' in doc.metadata:
                            filename = doc.metadata['uploaded_filename']
                            break
                    
                    if filename and hasattr(self, 'db_manager') and self.db_manager:
                        # Try to get file size from config or metadata
                        file_size = sum(len(doc.page_content) for doc in documents)
                        
                        success = self.db_manager.store_vector_points(
                            filename=filename,
                            point_ids=point_ids,
                            file_size=file_size,
                            collection_name=self.collection_name,
                            chunks_data=chunks_data
                        )
                        if success:
                            logger.info(f"ðŸ“Š Tracked {len(point_ids)} point IDs in database for {filename}")
                        else:
                            logger.warning(f"âš ï¸ Failed to track point IDs in database for {filename}")
                    else:
                        logger.warning("âš ï¸ No filename found in document metadata or database manager not available")
                except Exception as tracking_error:
                    logger.error(f"âŒ Failed to store vector tracking data: {tracking_error}")
                    
            except Exception as upload_error:
                logger.error(f"âŒ Failed to upload points to Qdrant: {upload_error}")
                raise
            
            logger.info(f"Successfully added {len(documents)} documents to vector store (total: {existing_count + len(documents)})")
            
        except Exception as e:
            logger.error(f"Error adding documents to vector store: {e}")
            raise
    
    def search(self, query: str, k: int = 3, score_threshold: float = 0.4) -> List[Document]:
        """Search for similar documents with enhanced filtering and ranking"""
        try:
            # Generate query embedding
            query_embedding = self.embeddings.embed_query(query)
            
            # Search in Qdrant with higher limit to allow for better filtering
            search_results = self.client.search(
                collection_name=self.collection_name,
                query_vector=query_embedding,
                limit=k * 3,  # Get more results for better filtering
                with_payload=True,
                score_threshold=max(0.3, score_threshold - 0.1)  # Slightly lower threshold for initial search
            )
            
            # Convert results to Document objects with enhanced filtering
            documents = []
            seen_content = set()  # Avoid duplicate content
            
            for result in search_results:
                # Skip results below the actual threshold
                if result.score < score_threshold:
                    continue
                
                # Try different content keys that might be stored
                content = (result.payload.get("content", "") or 
                          result.payload.get("text", "") or 
                          result.payload.get("processed_text", ""))
                
                # Skip empty or very short content
                if not content or len(content.strip()) < 50:
                    continue
                
                # Skip duplicate content (based on first 100 characters)
                content_hash = content[:100].strip().lower()
                if content_hash in seen_content:
                    continue
                seen_content.add(content_hash)
                
                # Remove content keys from metadata to avoid duplication
                metadata = {k: v for k, v in result.payload.items() 
                           if k not in ["content", "text", "processed_text"]}
                metadata["score"] = result.score
                
                # Add relevance boost for certain keywords
                query_lower = query.lower()
                content_lower = content.lower()
                relevance_boost = 0
                
                # Boost for exact keyword matches
                if any(word in content_lower for word in query_lower.split() if len(word) > 3):
                    relevance_boost += 0.1
                
                # Boost for banking-related terms
                banking_terms = ["bank", "Ð±Ð°Ð½Ðº", "Ñ„Ð¸Ð»Ð¸Ð°Ð»", "branch", "atm", "ÐºÐ°Ñ€Ñ‚Ð°", "card", "ÐºÑ€ÐµÐ´Ð¸Ñ‚", "credit"]
                if any(term in content_lower for term in banking_terms):
                    relevance_boost += 0.05
                
                metadata["adjusted_score"] = result.score + relevance_boost
                
                document = Document(
                    page_content=content,
                    metadata=metadata
                )
                documents.append(document)
                
                if len(documents) >= k:
                    break
            
            # Sort by adjusted score (highest first)
            documents.sort(key=lambda x: x.metadata.get("adjusted_score", x.metadata.get("score", 0)), reverse=True)
            
            logger.info(f"Retrieved {len(documents)} relevant documents for query: {query[:50]}...")
            return documents
            
        except Exception as e:
            logger.error(f"Error searching vector store: {e}")
            return []

    def delete_documents_by_source(self, source_identifier: str) -> int:
        """
        Enhanced document deletion using PostgreSQL point ID tracking for reliable cleanup
        
        This method uses the database tracking system to find exact point IDs,
        ensuring complete and accurate vector deletion.
        """
        try:
            logger.info(f"ðŸ—‘ï¸ Starting database-tracked vector deletion for: {source_identifier}")
            
            total_deleted = 0
            
            # METHOD 1: Use database tracking (primary method)
            if hasattr(self, 'db_manager') and self.db_manager:
                try:
                    # Get tracked point IDs from database
                    tracked_point_ids = self.db_manager.get_vector_points_for_file(
                        filename=source_identifier,
                        collection_name=self.collection_name
                    )
                    
                    if tracked_point_ids:
                        logger.info(f"ðŸ“Š Found {len(tracked_point_ids)} tracked point IDs for {source_identifier}")
                        
                        # Delete vectors using exact point IDs
                        try:
                            self.client.delete(
                                collection_name=self.collection_name,
                                points_selector=tracked_point_ids
                            )
                            
                            total_deleted = len(tracked_point_ids)
                            logger.info(f"âœ… Successfully deleted {total_deleted} vectors using database tracking")
                            
                            # Clean up database tracking records
                            deleted_records = self.db_manager.delete_vector_points_for_file(
                                filename=source_identifier,
                                collection_name=self.collection_name
                            )
                            logger.info(f"ðŸ§¹ Cleaned up {deleted_records} database tracking records")
                            
                        except Exception as delete_error:
                            logger.error(f"âŒ Failed to delete vectors from Qdrant: {delete_error}")
                            return 0
                            
                    else:
                        logger.warning(f"âš ï¸ No tracked point IDs found for {source_identifier}")
                        
                except Exception as db_error:
                    logger.error(f"âŒ Database tracking lookup failed: {db_error}")
                    
            else:
                logger.warning("âš ï¸ Database manager not available for tracking lookup")
            
            # METHOD 2: Fallback to metadata search (if database tracking failed)
            if total_deleted == 0:
                logger.info("ðŸ”„ Falling back to metadata search method...")
                
                # Enhanced source pattern matching for fallback cleanup
                source_patterns = [
                    source_identifier,                    # Direct filename match
                    f"uploaded_{source_identifier}",      # Standard upload prefix
                    f"documents/{source_identifier}",     # With directory path
                    f"./{source_identifier}",            # Relative path variant
                    source_identifier.replace("\\", "/"), # Normalize path separators
                    os.path.basename(source_identifier),  # Base filename only
                ]
                
                # Remove duplicates while preserving order
                unique_patterns = []
                for pattern in source_patterns:
                    if pattern not in unique_patterns:
                        unique_patterns.append(pattern)
                
                logger.info(f"ðŸ” Checking {len(unique_patterns)} source patterns for fallback cleanup")
                
                deletion_details = {}
                
                # Check each pattern and collect matching points
                for pattern in unique_patterns:
                    try:
                        filter_conditions = Filter(
                            should=[
                                FieldCondition(
                                    key="source",
                                    match=MatchValue(value=pattern)
                                ),
                                FieldCondition(
                                    key="uploaded_filename", 
                                    match=MatchValue(value=pattern)
                                ),
                                FieldCondition(
                                    key="filename",
                                    match=MatchValue(value=pattern)
                                )
                            ]
                        )
                        
                        # Scroll through all matching points
                        search_result = self.client.scroll(
                            collection_name=self.collection_name,
                            scroll_filter=filter_conditions,
                            limit=10000,  # Large limit to get all matching points
                            with_payload=True
                        )
                        
                        points_to_delete = []
                        for point in search_result[0]:  # scroll returns tuple (points, next_page_offset)
                            points_to_delete.append(point.id)
                            logger.debug(f"Found vector point {point.id} with payload: {point.payload}")
                        
                        if points_to_delete:
                            # Delete the points for this pattern
                            self.client.delete(
                                collection_name=self.collection_name,
                                points_selector=points_to_delete
                            )
                            
                            pattern_deleted = len(points_to_delete)
                            total_deleted += pattern_deleted
                            deletion_details[pattern] = pattern_deleted
                            
                            logger.info(f"âœ… Deleted {pattern_deleted} vectors for pattern: {pattern}")
                        
                    except Exception as pattern_error:
                        logger.warning(f"âš ï¸ Failed to process pattern '{pattern}': {pattern_error}")
                        continue
                
                # Log fallback results
                if total_deleted > 0:
                    logger.info(f"ðŸ”„ Fallback deletion completed: {total_deleted} vectors")
                    for pattern, count in deletion_details.items():
                        logger.info(f"   - Pattern '{pattern}': {count} vectors")
                else:
                    logger.warning(f"âŒ No vectors found using fallback method for: {source_identifier}")
            
            # Final verification (for both methods)
            if total_deleted > 0:
                logger.info(f"ðŸŽ¯ Successfully deleted {total_deleted} vectors for {source_identifier}")
            else:
                logger.warning(f"âŒ No vectors deleted for file: {source_identifier}")
                logger.info("ðŸ’¡ This might indicate the file was never processed or already deleted")
            
            return total_deleted
            
        except Exception as e:
            logger.error(f"âŒ Critical error during vector deletion: {e}")
            logger.exception("Full traceback:")
            return 0
    


def preprocess_documents(documents: List[Document]) -> List[Document]:
    """Enhanced document preprocessing for better retrieval quality"""
    processed_docs = []
    
    for doc in documents:
        # Clean content more thoroughly
        content = doc.page_content
        
        # Remove excessive whitespace and normalize
        content = re.sub(r'\s+', ' ', content)
        content = re.sub(r'\n\s*\n', '\n\n', content)  # Preserve paragraph breaks
        content = content.strip()
        
        # Remove common PDF artifacts
        content = re.sub(r'^\d+\s*$', '', content, flags=re.MULTILINE)  # Remove page numbers
        content = re.sub(r'^[^\w\s]*$', '', content, flags=re.MULTILINE)  # Remove lines with only symbols
        content = re.sub(r'\s*\.\s*\.\s*\.+', '...', content)  # Normalize ellipsis
        
        # Skip very short or low-quality content
        if len(content) < 100:
            continue
        
        # Skip content that's mostly numbers or symbols
        word_count = len(re.findall(r'\b\w+\b', content))
        if word_count < 10:
            continue
        
        # Enhance metadata
        metadata = doc.metadata.copy()
        metadata['word_count'] = word_count
        metadata['char_count'] = len(content)
        
        # Add content type hints
        content_lower = content.lower()
        if any(term in content_lower for term in ['Ñ‚ÐµÐ»ÐµÑ„Ð¾Ð½', 'phone', 'ÐºÐ¾Ð½Ñ‚Ð°ÐºÑ‚', 'contact', 'aloqa']):
            metadata['content_type'] = 'contact_info'
        elif any(term in content_lower for term in ['Ñ„Ð¸Ð»Ð¸Ð°Ð»', 'branch', 'Ð±Ð°Ð½ÐºÐ¾Ð¼Ð°Ñ‚', 'atm']):
            metadata['content_type'] = 'location_info'
        elif any(term in content_lower for term in ['Ð²Ñ€ÐµÐ¼Ñ', 'vaqt', 'time', 'Ð³Ñ€Ð°Ñ„Ð¸Ðº', 'schedule']):
            metadata['content_type'] = 'schedule_info'
        elif any(term in content_lower for term in ['ÑƒÑÐ»ÑƒÐ³Ð°', 'xizmat', 'service', 'ÐºÑ€ÐµÐ´Ð¸Ñ‚', 'kredit']):
            metadata['content_type'] = 'service_info'
        else:
            metadata['content_type'] = 'general_info'
        
        # Create processed document
        processed_doc = Document(
            page_content=content,
            metadata=metadata
        )
        processed_docs.append(processed_doc)
    
    logger.info(f"Preprocessed {len(processed_docs)} documents from {len(documents)} original documents")
    return processed_docs


async def process_uploaded_document(file: UploadFile, config: Dict[str, Any]) -> Tuple[List[Document], str]:
    """Ð®ÐºÐ»Ð°Ð½Ð³Ð°Ð½ Ò³ÑƒÐ¶Ð¶Ð°Ñ‚Ð½Ð¸ Ð¸ÑˆÐ»Ð°Ñˆ"""
    import tempfile
    import shutil
    
    # Calculate file size
    file.file.seek(0, 2)  # Seek to end
    file_size = file.file.tell()  # Get file size
    file.file.seek(0)  # Reset to beginning
    
    # Create temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{file.filename}") as temp_file:
        # Copy uploaded file content to temporary file
        shutil.copyfileobj(file.file, temp_file)
        temp_file_path = temp_file.name
    
    try:
        documents = []
        file_extension = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
        
        # Process based on file type
        if file_extension == 'pdf':
            loader = PyPDFLoader(temp_file_path)
            documents = loader.load()
            logger.info(f"Loaded PDF document: {file.filename}")
            
        elif file_extension in ['docx', 'doc']:
            loader = UnstructuredWordDocumentLoader(temp_file_path)
            documents = loader.load()
            logger.info(f"Loaded Word document: {file.filename}")
            
        elif file_extension in ['txt', 'md']:
            loader = TextLoader(temp_file_path)
            documents = loader.load()
            logger.info(f"Loaded text document: {file.filename}")
            
        else:
            # Try to load as text file
            try:
                loader = TextLoader(temp_file_path)
                documents = loader.load()
                logger.info(f"Loaded as text document: {file.filename}")
            except Exception as e:
                raise ValueError(f"Unsupported file type: {file_extension}. Supported types: PDF, DOCX, TXT, MD")
        
        if not documents:
            raise ValueError("No content could be extracted from the document")
        
        # Update metadata with upload information
        for doc in documents:
            doc.metadata.update({
                'uploaded_filename': file.filename,
                'upload_timestamp': datetime.now().isoformat(),
                'file_type': file_extension,
                'source': f"uploaded_{file.filename}"
            })
        
        # Preprocess documents
        processed_docs = preprocess_documents(documents)
        
        if not processed_docs:
            raise ValueError("Document processing resulted in no valid content")
        
        # Split documents into chunks - optimize for large documents
        # Use larger chunks for large documents to reduce processing time
        adaptive_chunk_size = config["chunk_size"]
        if file_size > 5 * 1024 * 1024:  # 5MB+
            adaptive_chunk_size = min(2000, config["chunk_size"] * 2)  # Double chunk size for large files
            logger.info(f"ðŸ“ˆ Using adaptive chunk size {adaptive_chunk_size} for large document")
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=adaptive_chunk_size,
            chunk_overlap=config["chunk_overlap"],
            length_function=len,
        )
        
        split_docs = text_splitter.split_documents(processed_docs)
        logger.info(f"Split uploaded document into {len(split_docs)} chunks")
        
        return split_docs, f"Successfully processed {file.filename}: {len(split_docs)} chunks created"
        
    finally:
        # Clean up temporary file
        try:
            os.unlink(temp_file_path)
        except Exception as e:
            logger.warning(f"Failed to delete temporary file {temp_file_path}: {e}")


async def process_uploaded_document_with_progress(file: UploadFile, config: Dict[str, Any], progress_callback=None) -> Tuple[List[Document], str]:
    """Process uploaded document with progress tracking"""
    
    if progress_callback:
        await progress_callback("validation", 5, "Ð¤Ð°Ð¹Ð»Ð½Ð¸ Ñ‚ÐµÐºÑˆÐ¸Ñ€Ð¸Ñˆ...")
    
    file_extension = Path(file.filename).suffix.lower()
    
    # Read file content
    content = await file.read()
    await file.seek(0)  # Reset file pointer
    
    if progress_callback:
        await progress_callback("reading", 15, "Ð¤Ð°Ð¹Ð» Ð¼Ð°Ð·Ð¼ÑƒÐ½Ð¸Ð½Ð¸ ÑžÒ›Ð¸Ñˆ...")
    
    # Process based on file type
    if file_extension == '.pdf':
        docs = await process_pdf_with_progress(content, file.filename, progress_callback)
    elif file_extension in ['.docx', '.doc']:
        docs = await process_docx_with_progress(content, file.filename, progress_callback)
    elif file_extension in ['.txt', '.md', '.py']:
        docs = await process_text_with_progress(content, file.filename, progress_callback)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")
    
    if progress_callback:
        await progress_callback("chunking", 70, "Ò²ÑƒÐ¶Ð¶Ð°Ñ‚Ð½Ð¸ Ð±ÑžÐ»Ð°ÐºÐ»Ð°Ñˆ...")
    
    # Split documents into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.get("chunk_size", 1000),
        chunk_overlap=config.get("chunk_overlap", 200)
    )
    
    split_docs = text_splitter.split_documents(docs)
    
    if progress_callback:
        await progress_callback("processing", 85, f"{len(split_docs)} Ñ‡Ð°Ð½Ðº Ñ‚Ð°Ð¹Ñ‘Ñ€")
    
    # Add uploaded filename to metadata
    for doc in split_docs:
        doc.metadata["uploaded_filename"] = file.filename
        doc.metadata["upload_timestamp"] = datetime.now().isoformat()
    
    message = f"Processed {len(split_docs)} chunks from uploaded document"
    
    if progress_callback:
        await progress_callback("completed", 100, "Ð˜ÑˆÐ»Ð¾Ð² Ð±ÐµÑ€Ð¸Ñˆ Ñ‚ÑƒÐ³Ð°Ñ‚Ð¸Ð»Ð´Ð¸")
    
    return split_docs, message


async def process_pdf_with_progress(content: bytes, filename: str, progress_callback=None) -> List[Document]:
    """Process PDF with progress tracking"""
    if progress_callback:
        await progress_callback("pdf_parsing", 25, "PDF Ñ„Ð°Ð¹Ð»Ð¸Ð½Ð¸ Ñ‚Ð°Ñ…Ð»Ð¸Ð» Ò›Ð¸Ð»Ð¸Ñˆ...")
    
    from PyPDF2 import PdfReader
    import io
    
    pdf_reader = PdfReader(io.BytesIO(content))
    documents = []
    
    total_pages = len(pdf_reader.pages)
    
    for i, page in enumerate(pdf_reader.pages):
        if progress_callback:
            page_progress = 25 + (i / total_pages) * 40  # Progress from 25% to 65%
            await progress_callback("pdf_pages", page_progress, f"PDF ÑÐ°Ò³Ð¸Ñ„Ð° {i+1}/{total_pages} Ð½Ð¸ ÑžÒ›Ð¸Ñˆ...")
        
        text = page.extract_text()
        if text.strip():
            doc = Document(
                page_content=text,
                metadata={
                    "source": f"uploaded_{filename}",
                    "page": i + 1,
                    "total_pages": total_pages
                }
            )
            documents.append(doc)
    
    return documents


async def process_docx_with_progress(content: bytes, filename: str, progress_callback=None) -> List[Document]:
    """Process DOCX with progress tracking"""
    if progress_callback:
        await progress_callback("docx_parsing", 30, "DOCX Ñ„Ð°Ð¹Ð»Ð¸Ð½Ð¸ Ñ‚Ð°Ñ…Ð»Ð¸Ð» Ò›Ð¸Ð»Ð¸Ñˆ...")
    
    from docx import Document as DocxDocument
    import io
    
    doc = DocxDocument(io.BytesIO(content))
    
    if progress_callback:
        await progress_callback("docx_text", 50, "DOCX Ð¼Ð°Ñ‚Ð½Ð¸Ð½Ð¸ Ñ‡Ð¸Ò›Ð°Ñ€Ð¸Ñˆ...")
    
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
    
    document = Document(
        page_content=text,
        metadata={
            "source": f"uploaded_{filename}",
            "paragraphs": len(doc.paragraphs)
        }
    )
    
    return [document]


async def process_text_with_progress(content: bytes, filename: str, progress_callback=None) -> List[Document]:
    """Process text files with progress tracking"""
    if progress_callback:
        await progress_callback("text_parsing", 40, "ÐœÐ°Ñ‚Ð½ Ñ„Ð°Ð¹Ð»Ð¸Ð½Ð¸ ÑžÒ›Ð¸Ñˆ...")
    
    try:
        text = content.decode('utf-8')
    except UnicodeDecodeError:
        try:
            text = content.decode('latin-1')
        except UnicodeDecodeError:
            text = content.decode('utf-8', errors='ignore')
    
    if progress_callback:
        await progress_callback("text_processed", 60, "ÐœÐ°Ñ‚Ð½ Ñ‚Ð°Ð¹Ñ‘Ñ€")
    
    document = Document(
        page_content=text,
        metadata={
            "source": f"uploaded_{filename}",
            "encoding": "utf-8",
            "lines": len(text.split('\n'))
        }
    )
    
    return [document]


def check_available_models(endpoint: str, api_key: str = "EMPTY") -> List[str]:
    """Check available models at vLLM endpoint"""
    try:
        import requests
        response = requests.get(
            f"{endpoint}/models",
            headers={"Authorization": f"Bearer {api_key}"},
            timeout=5
        )
        if response.status_code == 200:
            data = response.json()
            models = [model["id"] for model in data.get("data", [])]
            return models
        else:
            logger.warning(f"Failed to get models from {endpoint}: {response.status_code}")
            return []
    except Exception as e:
        logger.warning(f"Could not connect to {endpoint}: {e}")
        return []


async def init_vectorstore(config: Dict[str, Any], documents: List[Document], client: QdrantClient) -> SimpleQdrantVectorStore:
    """Initialize vector store with documents"""
    try:
        # Initialize embeddings
        global embeddings_model
        
        # Check if vLLM embedding endpoint is available
        vllm_endpoint = config.get("vllm_embedding_endpoint")
        if vllm_endpoint:
            available_models = check_available_models(vllm_endpoint, config["vllm_api_key"])
            if available_models:
                logger.info(f"Using vLLM embedding endpoint: {vllm_endpoint}")
                logger.info(f"Available models: {available_models}")
                
                # Use OpenAI-compatible embeddings with vLLM
                embeddings_model = OpenAIEmbeddings(
                    openai_api_base=vllm_endpoint,
                    openai_api_key=config["vllm_api_key"],
                    model=available_models[0] if available_models else "default"
                )
            else:
                logger.warning("vLLM embedding endpoint not available, falling back to local model")
                vllm_endpoint = None
        
        # Fallback to local embeddings
        if not vllm_endpoint:
            if HUGGINGFACE_AVAILABLE:
                logger.info("Using HuggingFace embeddings")
                embeddings_model = HuggingFaceEmbeddings(
                    model_name=config["embedding_model"],
                    model_kwargs={'device': 'cuda' if torch.cuda.is_available() else 'cpu'}
                )
            else:
                logger.info("Using custom multilingual E5 embeddings")
                embeddings_model = MultilingualE5Embeddings(
                    model_name=config["embedding_model"]
                )
        
        # Preprocess documents
        processed_docs = preprocess_documents(documents)
        logger.info(f"Preprocessed {len(processed_docs)} documents")
        
        # Create vector store with database manager for tracking
        vectorstore = SimpleQdrantVectorStore(
            client=client,
            collection_name=config["qdrant_collection_name"],
            embeddings=embeddings_model,
            config=config
        )
        
        # Inject database manager for vector tracking
        from database import get_db_manager
        vectorstore.db_manager = get_db_manager()
        
        # Add documents to vector store
        await vectorstore.add_documents(processed_docs)
        
        return vectorstore
        
    except Exception as e:
        logger.error(f"Error initializing vector store: {e}")
        raise


def init_llm(config: Dict[str, Any]) -> ChatOpenAI:
    """Initialize language model with enhanced parameters for detailed responses"""
    return ChatOpenAI(
        openai_api_base=config["vllm_chat_endpoint"],
        openai_api_key=config["vllm_api_key"],
        model_name=config["chat_model"],
        temperature=0.4,  # Balanced temperature for detailed but focused responses
        max_tokens=1024,  # Increased significantly for detailed responses
        request_timeout=45  # Increased timeout for longer response generation
    )


def get_qa_prompt() -> PromptTemplate:
    """
    Enhanced multilingual RAG QA prompt for accurate, context-aware responses.
    """

    template = """Siz Mohirdev kompaniyasidagi professional AI yordamchisiz. Quyidagi qoidalarga qat'iy amal qiling:

MUHIM: Faqat berilgan kontekst ma'lumotlariga asoslanib javob bering. Agar kontekstda javob yo'q bo'lsa, buni ochiq aytib bering.

1. KONTEKST TAHLILI:
   - Berilgan kontekst ma'lumotlarini diqqat bilan o'qing
   - Foydalanuvchi savoliga mos keladigan aniq ma'lumotlarni toping
   - Agar kontekstda javob yo'q bo'lsa, buni tan oling

2. JAVOB BERISH QOIDALARI:

   A) KONTEKST MAVJUD BO'LSA:
      - Faqat kontekstdagi ma'lumotlarga asoslanib javob bering
      - Aniq, batafsil va foydali javob bering
      - Kontekstdan misollar va raqamlar keltiring
      - Strukturali format ishlatilg (punktlar, ro'yxatlar)
      - Manba ma'lumotlarini kerak emas.
      - Gramatik to'g'ri, ma'noli va mazmunli javob bering.

   B) KONTEKST YETARLI BO'LMASA:
      - "Kechirasiz, bu savol bo'yicha hujjatlarda aniq ma'lumot topilmadi."
      - Qo'shimcha yordam taklif qiling
      - Boshqa mavzular bo'yicha yordam berish mumkinligini aytib bering

3. JAVOB FORMATI:
   - Aniq va tushunarli tilda yozing
   - 2-4 jumla yoki undan ko'proq (savolga qarab)
   - Amaliy ma'lumotlar va misollar qo'shing
   - Professional va do'stona ohang saqlang

4. Foydalanuvchi qaysi tilda yozsa, o'sha tilda javob bering.

5. Hech qachon [], (), {} belglarini ishlatmang.

6. Agar kontekstda ma'lumot bor bo'lsa, "Hi! I'm Mohirdev AI assistant" kabi umumiy javoblar bermang.

KONTEKST MA'LUMOTLARI:
{context}

FOYDALANUVCHI SAVOLI: {question}

ANIQ JAVOB (faqat kontekst asosida):"""

    return PromptTemplate(
        template=template,
        input_variables=["context", "question"]
    )

def format_docs(docs: List[Document]) -> str:
    """Format documents for context"""
    return "\n\n".join([doc.page_content for doc in docs])


def clean_response(response: str) -> str:
    """Clean up LLM response by removing unwanted markers and formatting"""
    import re
    
    # Remove reference markers like [[1]], [[doc_1]], [[none]], etc.
    response = re.sub(r'\[\[.*?\]\]', '', response)
    
    # Remove other common markers
    response = re.sub(r'\[.*?\]', '', response)
    
    # Remove extra whitespace
    response = re.sub(r'\s+', ' ', response)
    response = response.strip()
    
    return response


def create_qa_chain(vectorstore: SimpleQdrantVectorStore, llm: ChatOpenAI, prompt: PromptTemplate, config: Dict[str, Any]):
    """Create QA chain with improved accuracy and context filtering"""
    
    def qa_function(question: str, chat_history: List[ChatMessage] = None) -> str:
        try:
            # Check if this is a simple greeting
            simple_greetings = ["salom", "hello", "hi", "Ð¿Ñ€Ð¸Ð²ÐµÑ‚", "assalomu alaykum", "good morning", "good afternoon"]
            if question.lower().strip() in simple_greetings:
                if "salom" in question.lower() or "assalomu" in question.lower():
                    return "Salom! Men Mohirdev AI yordamchisi. Sizga qanday yordam bera olaman?"
                elif "Ð¿Ñ€Ð¸Ð²ÐµÑ‚" in question.lower():
                    return "ÐŸÑ€Ð¸Ð²ÐµÑ‚! Ð¯ AI-Ð¿Ð¾Ð¼Ð¾Ñ‰Ð½Ð¸Ðº Mohirdev. Ð§ÐµÐ¼ Ð¼Ð¾Ð³Ñƒ Ð¿Ð¾Ð¼Ð¾Ñ‡ÑŒ?"
                else:
                    return "Hi! I'm Mohirdev AI assistant. How can I help you?"
            
            # Retrieve relevant documents with improved filtering
            # Use higher score threshold for better quality
            docs = vectorstore.search(question, k=config.get("top_k", 5), score_threshold=0.6)
            
            # If no high-quality docs found, try with lower threshold
            if not docs:
                docs = vectorstore.search(question, k=config.get("top_k", 3), score_threshold=0.4)
            
            # Format context with relevant documents - enhanced filtering
            if docs:
                # Filter out very short or low-quality content
                quality_docs = []
                for doc in docs:
                    content = doc.page_content.strip()
                    score = doc.metadata.get('score', 0)
                    
                    # Skip very short content or very low scores
                    if len(content) < 100 or score < 0.4:
                        continue
                    
                    quality_docs.append(doc)
                
                if quality_docs:
                    context_parts = []
                    for i, doc in enumerate(quality_docs, 1):
                        content = doc.page_content.strip()
                        score = doc.metadata.get('score', 0)
                        # Add document with better formatting
                        context_parts.append(f"Ma'lumot {i} (relevantlik: {score:.2f}):\n{content}")
                    
                    context = "\n\n".join(context_parts)
                    # Log the scores for debugging
                    scores = [doc.metadata.get('score', 0) for doc in quality_docs]
                    logger.info(f"Using {len(quality_docs)} quality documents with scores: {scores}")
                else:
                    context = "Hech qanday tegishli hujjat topilmadi."
                    logger.warning(f"No quality documents found for query: {question}")
            else:
                context = "Hech qanday tegishli hujjat topilmadi."
                logger.warning(f"No documents found for query: {question}")
            
            # Enhanced prompt formatting with better context handling
            if context and context != "Hech qanday tegishli hujjat topilmadi.":
                # Use context-aware prompt
                formatted_prompt = prompt.template.replace("{context}", context).replace("{question}", question)
            else:
                # Use fallback prompt for when no relevant context is found
                fallback_template = """Siz Mohirdev kompaniyasidagi professional AI yordamchisiz. 

Foydalanuvchi savoli: {question}

Kechirasiz, bu savol bo'yicha aniq ma'lumot topilmadi. Iltimos, savolingizni aniqroq qo'ying yoki boshqa mavzu bo'yicha so'rang. 

Men quyidagi mavzularda yordam bera olaman:
- Bank xizmatlari haqida umumiy ma'lumotlar
- Aloqa ma'lumotlari
- Filiallar va ATM-lar haqida
- Internet-banking xizmatlari
- Boshqa bank xizmatlari

Qanday yordam kerak?"""
                formatted_prompt = fallback_template.replace("{question}", question)
            
            logger.debug(f"Formatted prompt length: {len(formatted_prompt)}")
            
            # Use messages format for ChatOpenAI
            from langchain.schema import HumanMessage
            messages = [HumanMessage(content=formatted_prompt)]
            
            response = llm.invoke(messages)
            logger.debug(f"LLM response type: {type(response)}")
            logger.info(f"Generated response length: {len(str(response)) if response else 0}")
            
            # Handle different response types
            if hasattr(response, 'content'):
                response_text = response.content
            elif hasattr(response, 'text'):
                response_text = response.text
            elif isinstance(response, str):
                response_text = response
            elif isinstance(response, dict):
                response_text = response.get('content', response.get('text', str(response)))
            else:
                response_text = str(response)
            
            # Clean and validate response
            cleaned_response = clean_response(response_text)
            
            # Ensure response is not empty or too generic
            if not cleaned_response or len(cleaned_response.strip()) < 10:
                return "Kechirasiz, javob olishda xato yuz berdi. Iltimos, savolingizni boshqacha qo'yib ko'ring."
            
            return cleaned_response
            
        except Exception as e:
            logger.error(f"Error generating response: {e}")
            logger.error(f"Error type: {type(e)}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return "Kechirasiz, javob olishda xato yuz berdi. Iltimos, qaytadan urinib ko'ring."
    
    return qa_function


async def initialize_rag_system():
    """Initialize RAG system components"""
    global vectorstore, qa_chain, qdrant_client, db_manager
    
    try:
        logger.info("Initializing RAG system...")
        
        # Load configuration
        config = load_config()
        
        # Initialize database
        logger.info("Initializing database...")
        db_manager = DatabaseManager(config["database_url"])
        
        # Migrate from JSON if exists
        if os.path.exists(CHAT_HISTORY_FILE):
            logger.info("Migrating chat history from JSON to PostgreSQL...")
            db_manager.migrate_from_json(CHAT_HISTORY_FILE)
        
        # Initialize Qdrant client
        logger.info("Initializing Qdrant client...")
        qdrant_client = init_qdrant_client(config)
        
        # Load and process documents
        logger.info("Loading documents...")
        documents = load_and_split_documents(config)
        
        # Initialize vector store
        logger.info("Initializing vector store...")
        vectorstore = await init_vectorstore(config, documents, qdrant_client)
        
        # Initialize LLM
        logger.info("Initializing language model...")
        llm = init_llm(config)
        
        # Create QA chain
        logger.info("Creating QA chain...")
        prompt = get_qa_prompt()
        qa_chain = create_qa_chain(vectorstore, llm, prompt, config)
        
        logger.info("RAG system initialized successfully!")
        
    except Exception as e:
        logger.error(f"Failed to initialize RAG system: {e}")
        raise


# FastAPI app initialization
@asynccontextmanager
async def lifespan(app: FastAPI):
    """Application lifespan manager"""
    # Startup
    await initialize_rag_system()
    yield
    # Shutdown
    if db_manager:
        db_manager.close()
    logger.info("Application shutdown complete")


app = FastAPI(
    title="Platform Assistant RAG API",
    description="RAG-based chatbot API with PostgreSQL database",
    version="2.0.0",
    lifespan=lifespan
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health", response_model=HealthResponse)
async def health_check():
    """Health check endpoint"""
    try:
        # Check Qdrant
        qdrant_status = "healthy"
        try:
            if qdrant_client:
                collections = qdrant_client.get_collections()
                qdrant_status = f"healthy ({len(collections.collections)} collections)"
            else:
                qdrant_status = "not initialized"
        except Exception as e:
            qdrant_status = f"error: {str(e)}"
        
        # Check database
        database_status = "healthy"
        try:
            if db_manager:
                with db_manager.get_connection() as conn:
                    with conn.cursor() as cur:
                        cur.execute("SELECT 1")
                database_status = "healthy"
            else:
                database_status = "not initialized"
        except Exception as e:
            database_status = f"error: {str(e)}"
        
        return HealthResponse(
            status="healthy",
            message="RAG system is operational",
            qdrant_status=qdrant_status,
            database_status=database_status
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Health check failed: {str(e)}")


@app.post("/v1/chat/completions", response_model=ChatCompletionResponse)
async def chat_completions(request: ChatCompletionRequest):
    """OpenAI-compatible chat completions endpoint"""
    try:
        if not qa_chain:
            raise HTTPException(status_code=503, detail="RAG system not initialized")
        
        # Get the last user message
        user_messages = [msg for msg in request.messages if msg.role == "user"]
        if not user_messages:
            raise HTTPException(status_code=400, detail="No user message found")
        
        last_message = user_messages[-1].content
        
        # Generate response using RAG
        response_text = qa_chain(last_message, request.messages)
        
        # Format response in OpenAI format
        completion_id = f"chatcmpl-{uuid.uuid4().hex[:8]}"
        
        return ChatCompletionResponse(
            id=completion_id,
            created=int(time.time()),
            model=request.model,
            choices=[{
                "index": 0,
                "message": {
                    "role": "assistant",
                    "content": response_text
                },
                "finish_reason": "stop"
            }],
            usage={
                "prompt_tokens": len(last_message.split()),
                "completion_tokens": len(response_text.split()),
                "total_tokens": len(last_message.split()) + len(response_text.split())
            }
        )
        
    except Exception as e:
        logger.error(f"Chat completion error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/v1/collections")
async def get_collections():
    """Get Qdrant collections info"""
    try:
        if not qdrant_client:
            raise HTTPException(status_code=503, detail="Qdrant client not initialized")
        
        collections = qdrant_client.get_collections()
        
        collection_info = []
        for collection in collections.collections:
            try:
                info = qdrant_client.get_collection(collection.name)
                collection_info.append({
                    "name": collection.name,
                    "points_count": info.points_count,
                    "vectors_count": info.vectors_count,
                    "status": info.status
                })
            except Exception as e:
                collection_info.append({
                    "name": collection.name,
                    "error": str(e)
                })
        
        return {
            "collections": collection_info,
            "total_collections": len(collections.collections)
        }
        
    except Exception as e:
        logger.error(f"Collections error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/v1/chat/new", response_model=NewChatResponse)
async def create_new_chat(request: NewChatRequest):
    """Create a new chat session for user (New Chat button functionality)"""
    try:
        if not db_manager:
            raise HTTPException(status_code=503, detail="Database not initialized")
        
        # Create new chat session
        new_chat_id = db_manager.create_new_chat_for_user(request.user_id)
        if not new_chat_id:
            raise HTTPException(status_code=500, detail="Failed to create new chat session")
        
        # Get the created session
        session = db_manager.get_chat_session(new_chat_id)
        if not session:
            raise HTTPException(status_code=500, detail="Failed to retrieve created session")
        
        return NewChatResponse(
            chat_id=new_chat_id,
            user_id=request.user_id,
            message="Ð¯Ð½Ð³Ð¸ Ñ‡Ð°Ñ‚ ÑÑ€Ð°Ñ‚Ð¸Ð»Ð´Ð¸. ÒšÐ°Ð½Ð´Ð°Ð¹ ÑÐ°Ð²Ð¾Ð»Ð¸Ð½Ð³Ð¸Ð· Ð±Ð¾Ñ€?",
            created_at=session["created_at"],
            last_activity=session["last_activity"]
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"New chat creation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/v1/user/session-status", response_model=UserSessionStatusResponse)
async def get_user_session_status(request: UserSessionStatusRequest):
    """Check user session status and activity"""
    try:
        if not db_manager:
            raise HTTPException(status_code=503, detail="Database not initialized")
        
        # Check for active session
        active_session = db_manager.get_active_user_session(request.user_id)
        
        if active_session:
            return UserSessionStatusResponse(
                user_id=request.user_id,
                has_active_session=True,
                active_chat_id=active_session["chat_id"],
                last_activity=active_session["last_activity"],
                session_expired=False
            )
        else:
            # Check if user has any sessions (but expired)
            user_last_activity = db_manager.get_user_last_activity(request.user_id)
            
            return UserSessionStatusResponse(
                user_id=request.user_id,
                has_active_session=False,
                active_chat_id=None,
                last_activity=user_last_activity.isoformat() if user_last_activity else None,
                session_expired=user_last_activity is not None
            )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Session status error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/v1/chat", response_model=ChatResponse)
async def chat_with_history(request: ChatRequest):
    """Chat with history tracking using PostgreSQL and time-based session management"""
    try:
        if not qa_chain or not db_manager:
            raise HTTPException(status_code=503, detail="System not initialized")
        
        # Get or create chat session with improved time-based management
        session = db_manager.get_or_create_session(request.user_id, request.chat_id)
        if not session:
            raise HTTPException(status_code=500, detail="Failed to create chat session")
        
        chat_id = session["chat_id"]
        
        # Log session activity for debugging
        logger.debug(f"Using chat session {chat_id} for user {request.user_id}")
        
        # Add user message to database
        if not db_manager.add_message(chat_id, "user", request.message):
            raise HTTPException(status_code=500, detail="Failed to save user message")
        
        # Get chat history for context
        messages = db_manager.get_chat_messages(chat_id)
        chat_history = [ChatMessage(**msg) for msg in messages]
        
        # Generate response using RAG with improved embedding filtering
        response_text = qa_chain(request.message, chat_history)
        
        # Add assistant response to database
        if not db_manager.add_message(chat_id, "assistant", response_text):
            logger.warning("Failed to save assistant message")
        
        return ChatResponse(
            chat_id=chat_id,
            user_id=request.user_id,
            message=response_text,
            timestamp=datetime.now()
        )
        
    except Exception as e:
        logger.error(f"Chat error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/v1/chat/history")
async def get_chat_history(request: ChatHistoryRequest):
    """Get chat history from PostgreSQL"""
    try:
        if not db_manager:
            raise HTTPException(status_code=503, detail="Database not initialized")
        
        # Get chat session
        session = db_manager.get_chat_session(request.chat_id)
        if not session:
            raise HTTPException(status_code=404, detail="Chat session not found")
        
        # Verify user access
        if session["user_id"] != request.user_id:
            raise HTTPException(status_code=403, detail="Access denied")
        
        messages = session["messages"]
        
        return {
            "chat_id": request.chat_id,
            "user_id": request.user_id,
            "messages": messages,
            "total_messages": len(messages),
            "created_at": session["created_at"],
            "last_activity": session["last_activity"]
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Chat history error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/v1/user/{user_id}/chats")
async def get_user_chats(user_id: str):
    """Get all chats for a user from PostgreSQL"""
    try:
        if not db_manager:
            raise HTTPException(status_code=503, detail="Database not initialized")
        
        chats = db_manager.get_user_chats(user_id)
        
        return {
            "user_id": user_id,
            "chats": chats,
            "total_chats": len(chats)
        }
        
    except Exception as e:
        logger.error(f"User chats error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Global progress storage for tracking upload progress
upload_progress_store = {}

@app.post("/v1/documents/upload", response_model=DocumentUploadResponse)
async def upload_document(file: UploadFile = File(...)):
    """Ò²ÑƒÐ¶Ð¶Ð°Ñ‚ ÑŽÐºÐ»Ð°Ñˆ Ð²Ð° Ð²ÐµÐºÑ‚Ð¾Ñ€Ð»Ð°ÑˆÑ‚Ð¸Ñ€Ð¸Ñˆ (progress trackingsiz)"""
    return await _upload_document_internal(file, None)

@app.post("/v1/documents/upload-with-progress")
async def upload_document_with_progress(file: UploadFile = File(...)):
    """Ò²ÑƒÐ¶Ð¶Ð°Ñ‚ ÑŽÐºÐ»Ð°Ñˆ Ð²Ð° Ð²ÐµÐºÑ‚Ð¾Ñ€Ð»Ð°ÑˆÑ‚Ð¸Ñ€Ð¸Ñˆ (progress tracking bilan)"""
    # Generate unique upload ID
    upload_id = str(uuid.uuid4())
    
    # Initialize progress
    upload_progress_store[upload_id] = {
        "stage": "starting",
        "progress": 0,
        "message": "Starting upload process...",
        "filename": file.filename
    }
    
    # Start upload in background
    asyncio.create_task(_upload_with_progress_task(file, upload_id))
    
    return {"upload_id": upload_id, "message": "Upload started, use /v1/documents/upload-progress/{upload_id} for progress"}

@app.get("/v1/documents/upload-progress/{upload_id}")
async def get_upload_progress(upload_id: str):
    """Get upload progress using Server-Sent Events"""
    
    async def generate_progress():
        while upload_id in upload_progress_store:
            progress_data = upload_progress_store[upload_id]
            yield f"data: {json.dumps(progress_data)}\n\n"
            
            # If completed or failed, send final message and exit
            if progress_data.get("progress", 0) >= 100 or progress_data.get("stage") in ["completed", "failed"]:
                # Clean up after 5 seconds
                await asyncio.sleep(5)
                if upload_id in upload_progress_store:
                    del upload_progress_store[upload_id]
                break
                
            await asyncio.sleep(0.5)  # Update every 500ms
    
    return StreamingResponse(
        generate_progress(),
        media_type="text/plain",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Content-Type": "text/event-stream"
        }
    )

async def _upload_with_progress_task(file: UploadFile, upload_id: str):
    """Background task for upload with progress"""
    try:
        async def progress_callback(stage: str, progress: float, message: str):
            upload_progress_store[upload_id] = {
                "stage": stage,
                "progress": min(progress, 100),
                "message": message,
                "filename": file.filename
            }
        
        result = await _upload_document_internal(file, progress_callback)
        
        # Final success message
        upload_progress_store[upload_id] = {
            "stage": "completed",
            "progress": 100,
            "message": f"Upload completed successfully! {result.chunks_added} chunks processed.",
            "filename": file.filename,
            "result": result.dict()
        }
        
    except Exception as e:
        # Final error message
        upload_progress_store[upload_id] = {
            "stage": "failed",
            "progress": 0,
            "message": f"Upload failed: {str(e)}",
            "filename": file.filename,
            "error": str(e)
        }

async def _upload_document_internal(file: UploadFile, progress_callback=None):
    """Ò²ÑƒÐ¶Ð¶Ð°Ñ‚ ÑŽÐºÐ»Ð°Ñˆ Ð²Ð° Ð²ÐµÐºÑ‚Ð¾Ñ€Ð»Ð°ÑˆÑ‚Ð¸Ñ€Ð¸Ñˆ"""
    start_time = time.time()
    file_size = 0  # Initialize file_size early to avoid NameError
    
    try:
        if not vectorstore or not qdrant_client:
            raise HTTPException(status_code=503, detail="Vector store not initialized")
        
        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        # Check file size (limit to 50MB)
        content = await file.read()
        file_size = len(content)
        
        if file_size > 50 * 1024 * 1024:  # 50MB limit
            raise HTTPException(status_code=400, detail="File too large. Maximum size is 50MB")
        
        # Warn about large files that might cause timeouts
        if file_size > 5 * 1024 * 1024:  # 5MB warning
            logger.warning(f"âš ï¸ Large file detected ({file_size / (1024*1024):.1f}MB). Processing may take 2-5 minutes.")
        
        if progress_callback:
            await progress_callback("validation", 5, f"File validated ({file_size / (1024*1024):.1f}MB)")
        
        # Reset file pointer
        await file.seek(0)
        
        # Get config
        config = load_config()
        
        if progress_callback:
            await progress_callback("processing", 10, "Processing document content...")
        
        # Process the uploaded document
        split_docs, process_message = await process_uploaded_document(file, config)
        
        if progress_callback:
            await progress_callback("processed", 25, f"Document processed: {len(split_docs)} chunks created")
        
        # Debug: Log what we're about to add
        logger.info(f"ðŸ” About to add {len(split_docs)} documents to vector store")
        for i, doc in enumerate(split_docs[:3]):  # Log first 3 documents
            logger.info(f"ðŸ“„ Document {i+1}: {len(doc.page_content)} chars, metadata: {doc.metadata}")
        
        # Add documents to vector store
        if progress_callback:
            await progress_callback("embedding", 30, "Starting vectorization process...")
        
        await vectorstore.add_documents(split_docs, progress_callback=progress_callback)
        
        if progress_callback:
            await progress_callback("vectorized", 80, "Vectorization completed, saving file...")
        
        # Save file to documents directory
        documents_path = Path(config["documents_path"])
        documents_path.mkdir(exist_ok=True)
        
        file_path = documents_path / file.filename
        with open(file_path, "wb") as f:
            await file.seek(0)
            content = await file.read()
            f.write(content)
        
        processing_time = time.time() - start_time
        
        if progress_callback:
            await progress_callback("saving", 90, "Saving file to storage...")
        
        logger.info(f"Successfully uploaded and processed document: {file.filename}")
        
        # Calculate performance metrics
        chunks_per_second = len(split_docs) / max(processing_time, 0.1)
        mb_per_second = (file_size / (1024*1024)) / max(processing_time, 0.1)
        
        performance_msg = f"âš¡ Performance: {chunks_per_second:.1f} chunks/sec, {mb_per_second:.2f} MB/sec"
        logger.info(performance_msg)
        
        if progress_callback:
            await progress_callback("completed", 100, f"Upload completed! {performance_msg}")
        
        return DocumentUploadResponse(
            success=True,
            message=f"Document successfully uploaded and vectorized. {process_message} {performance_msg}",
            filename=file.filename,
            file_size=file_size,
            chunks_added=len(split_docs),
            processing_time=round(processing_time, 2)
        )
        
    except HTTPException:
        raise
    except Exception as e:
        processing_time = time.time() - start_time
        logger.error(f"Document upload error: {e}")
        
        return DocumentUploadResponse(
            success=False,
            message=f"Failed to process document: {str(e)}",
            filename=file.filename if file.filename else "unknown",
            file_size=file_size,
            chunks_added=0,
            processing_time=round(processing_time, 2)
        )


@app.get("/v1/documents/list", response_model=DocumentListResponse)
async def list_documents():
    """Ò²ÑƒÐ¶Ð¶Ð°Ñ‚Ð»Ð°Ñ€ Ñ€ÑžÐ¹Ñ…Ð°Ñ‚Ð¸Ð½Ð¸ Ð¾Ð»Ð¸Ñˆ"""
    try:
        config = load_config()
        documents_path = Path(config["documents_path"])
        
        if not documents_path.exists():
            return DocumentListResponse(
                success=True,
                message="Documents directory does not exist",
                files=[],
                total_files=0,
                total_size=0
            )
        
        files_info = []
        total_size = 0
        
        # Supported file extensions
        supported_extensions = {'.pdf', '.docx', '.doc', '.txt', '.md', '.py'}
        
        for file_path in documents_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                stat = file_path.stat()
                file_size = stat.st_size
                total_size += file_size
                
                files_info.append(FileInfo(
                    filename=file_path.name,
                    file_size=file_size,
                    created_at=datetime.fromtimestamp(stat.st_ctime).isoformat(),
                    modified_at=datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    file_extension=file_path.suffix.lower()
                ))
        
        # Sort files by modified time (newest first)
        files_info.sort(key=lambda x: x.modified_at, reverse=True)
        
        return DocumentListResponse(
            success=True,
            message=f"Found {len(files_info)} documents",
            files=files_info,
            total_files=len(files_info),
            total_size=total_size
        )
        
    except Exception as e:
        logger.error(f"Document list error: {e}")
        return DocumentListResponse(
            success=False,
            message=f"Failed to list documents: {str(e)}",
            files=[],
            total_files=0,
            total_size=0
        )


@app.delete("/v1/documents/delete", response_model=DocumentDeleteResponse)
async def delete_document(request: DocumentDeleteRequest):
    """Ò²ÑƒÐ¶Ð¶Ð°Ñ‚Ð½Ð¸ ÑžÑ‡Ð¸Ñ€Ð¸Ñˆ Ð²Ð° ÑƒÐ½Ð¸Ð½Ð³ Ð²ÐµÐºÑ‚Ð¾Ñ€Ð»Ð°Ñ€Ð¸Ð½Ð¸ Ñ‚Ð¾Ð·Ð°Ð»Ð°Ñˆ"""
    try:
        if not vectorstore or not qdrant_client:
            raise HTTPException(status_code=503, detail="Vector store not initialized")
        
        config = load_config()
        documents_path = Path(config["documents_path"])
        file_path = documents_path / request.filename
        
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="File not found")
        
        if not file_path.is_file():
            raise HTTPException(status_code=400, detail="Path is not a file")
        
        # Check if file is in documents directory (security check)
        if not str(file_path.resolve()).startswith(str(documents_path.resolve())):
            raise HTTPException(status_code=403, detail="Access denied")
        
        # Delete embeddings from vector store first using enhanced deletion
        embeddings_deleted = 0
        try:
            # Use the enhanced vector deletion method
            embeddings_deleted = vectorstore.delete_documents_by_source(request.filename)
            logger.info(f"ðŸ—‘ï¸ Enhanced vector deletion completed: {embeddings_deleted} embeddings removed")
            
        except Exception as e:
            logger.warning(f"âš ï¸ Vector deletion failed for {request.filename}: {e}")
            # Continue with file deletion even if embeddings deletion fails
        
        # Delete the file from filesystem
        file_path.unlink()
        
        logger.info(f"Successfully deleted document: {request.filename} (with {embeddings_deleted} embeddings)")
        
        message = f"Document '{request.filename}' successfully deleted"
        if embeddings_deleted > 0:
            message += f" along with {embeddings_deleted} embeddings"
        else:
            message += " (no embeddings found to delete)"
        
        return DocumentDeleteResponse(
            success=True,
            message=message,
            filename=request.filename,
            embeddings_deleted=embeddings_deleted
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document delete error: {e}")
        return DocumentDeleteResponse(
            success=False,
            message=f"Failed to delete document: {str(e)}",
            filename=request.filename,
            embeddings_deleted=0
        )





@app.exception_handler(Exception)
async def global_exception_handler(request, exc):
    """Global exception handler"""
    logger.error(f"Unhandled exception: {exc}")
    return JSONResponse(
        status_code=500,
        content={"detail": "Internal server error"}
    )


if __name__ == "__main__":
    import uvicorn
    
    config = load_config()
    
    logger.info(f"Starting server on {config['host']}:{config['port']}")
    uvicorn.run(
        "app_postgres:app",
        host=config["host"],
        port=config["port"],
        reload=False,
        timeout_keep_alive=300,  # Keep connections alive for 5 minutes
        timeout_graceful_shutdown=30  # Allow 30 seconds for graceful shutdown
    ) 